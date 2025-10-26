"""
OCR Text Extraction Module for ProctorIQ
Handles extraction of text from images and PDF files
"""

import os
import io
from typing import Union, List
from pathlib import Path
import logging

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logging.warning("Tesseract OCR not available. Install pytesseract and pillow for OCR functionality.")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available. Install PyMuPDF for PDF text extraction.")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install python-docx for Word document support.")

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False

class OCRExtractor:
    """OCR text extraction from images and PDFs"""
    
    def __init__(self):
        """Initialize OCR extractor"""
        self.supported_image_formats = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'}
        self.supported_pdf_formats = {'.pdf'}
        self.supported_text_formats = {'.txt', '.md', '.csv'}
        self.supported_doc_formats = {'.docx', '.doc', '.rtf'}
        # Add support for code and programming files
        self.supported_code_formats = {
            '.py', '.js', '.java', '.cpp', '.c', '.h', '.hpp', 
            '.cs', '.rb', '.php', '.swift', '.go', '.rs', '.ts', 
            '.jsx', '.tsx', '.json', '.xml', '.html', '.css', '.scss',
            '.sh', '.bash', '.sql', '.r', '.m', '.kt', '.scala'
        }
        
        # Configure Tesseract if available
        if TESSERACT_AVAILABLE:
            self._configure_tesseract()
    
    def _configure_tesseract(self):
        """Configure Tesseract OCR settings"""
        try:
            # Try to find Tesseract executable on macOS
            possible_paths = [
                '/usr/local/bin/tesseract',
                '/opt/homebrew/bin/tesseract',
                '/usr/bin/tesseract'
            ]
            
            for path in possible_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break
        except Exception as e:
            logging.warning(f"Tesseract configuration warning: {e}")
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from image or PDF file
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        # Check for code files first (treat as text)
        if file_extension in self.supported_code_formats:
            return self._extract_from_text_file(file_path)
        elif file_extension in self.supported_image_formats:
            return self._extract_from_image(file_path)
        elif file_extension in self.supported_pdf_formats:
            return self._extract_from_pdf(file_path)
        elif file_extension in self.supported_text_formats:
            return self._extract_from_text_file(file_path)
        elif file_extension in self.supported_doc_formats:
            return self._extract_from_document(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_text_file(self, text_path: Path) -> str:
        """
        Extract text from text files (TXT, MD, CSV)
        
        Args:
            text_path: Path to text file
            
        Returns:
            File content as text
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(text_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    # If we successfully read the file, process based on extension
                    file_extension = text_path.suffix.lower()
                    
                    if file_extension == '.csv':
                        # For CSV files, format as readable text
                        return self._format_csv_content(content)
                    else:
                        # For TXT, MD files, return as-is
                        return self._clean_text(content)
                        
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, return error message
            return f"[TEXT EXTRACTION FAILED] Could not decode file {text_path.name} with any supported encoding"
            
        except Exception as e:
            logging.error(f"Text extraction failed for {text_path}: {e}")
            return f"[TEXT EXTRACTION FAILED] Error reading {text_path.name}: {str(e)}"
    
    def _extract_from_document(self, doc_path: Path) -> str:
        """
        Extract text from Word documents and RTF files
        
        Args:
            doc_path: Path to document file
            
        Returns:
            Extracted text content
        """
        file_extension = doc_path.suffix.lower()
        
        if file_extension == '.docx':
            return self._extract_from_docx(doc_path)
        elif file_extension in ['.doc', '.rtf']:
            return f"[DOCUMENT FORMAT NOT SUPPORTED] {file_extension.upper()} files require conversion to DOCX or text format"
        else:
            return f"[UNSUPPORTED DOCUMENT FORMAT] {file_extension}"
    
    def _extract_from_docx(self, docx_path: Path) -> str:
        """
        Extract text from DOCX files
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        if not PYTHON_DOCX_AVAILABLE:
            return f"[DOCX EXTRACTION NOT AVAILABLE] Install python-docx package for Word document support"
        
        try:
            doc = Document(docx_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return self._clean_text("\n".join(text_content))
            
        except Exception as e:
            logging.error(f"DOCX extraction failed for {docx_path}: {e}")
            return f"[DOCX EXTRACTION FAILED] Error extracting text from {docx_path.name}: {str(e)}"
    
    def _format_csv_content(self, csv_content: str) -> str:
        """
        Format CSV content as readable text
        
        Args:
            csv_content: Raw CSV content
            
        Returns:
            Formatted text
        """
        try:
            import io
            import csv as csv_module
            
            lines = []
            csv_reader = csv_module.reader(io.StringIO(csv_content))
            
            for row_num, row in enumerate(csv_reader, 1):
                if row:  # Skip empty rows
                    formatted_row = f"Row {row_num}: " + " | ".join(row)
                    lines.append(formatted_row)
            
            return "\n".join(lines)
            
        except Exception as e:
            logging.error(f"CSV formatting failed: {e}")
            return csv_content  # Return raw content if formatting fails
    
    def _extract_from_image(self, image_path: Path) -> str:
        """
        Extract text from image using OCR
        
        Args:
            image_path: Path to image file
            
        Returns:
            Extracted text
        """
        if not TESSERACT_AVAILABLE:
            return f"[OCR NOT AVAILABLE] Image file: {image_path.name}"
        
        try:
            # Open and process image
            with Image.open(image_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Extract text using Tesseract
                text = pytesseract.image_to_string(
                    img,
                    config='--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,?!()[]{}:;-_+=*/\'"@#$%^&|\\~` \n\t'
                )
                
                # Clean up extracted text
                return self._clean_text(text)
                
        except Exception as e:
            logging.error(f"OCR extraction failed for {image_path}: {e}")
            return f"[OCR FAILED] Error extracting text from {image_path.name}: {str(e)}"
    
    def _extract_from_pdf(self, pdf_path: Path) -> str:
        """
        Extract text from PDF file
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text from all pages
        """
        if not PYMUPDF_AVAILABLE:
            return f"[PDF EXTRACTION NOT AVAILABLE] PDF file: {pdf_path.name}"
        
        try:
            text_content = []
            
            # Open PDF document
            with fitz.open(pdf_path) as doc:
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Try text extraction first
                    text = page.get_text()
                    
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                    else:
                        # If no text found, try OCR on page image
                        if TESSERACT_AVAILABLE:
                            try:
                                # Get page as image
                                mat = fitz.Matrix(2.0, 2.0)  # Increase resolution
                                pix = page.get_pixmap(matrix=mat)
                                img_data = pix.tobytes("png")
                                
                                # Convert to PIL Image and extract text
                                img = Image.open(io.BytesIO(img_data))
                                ocr_text = pytesseract.image_to_string(img)
                                
                                if ocr_text.strip():
                                    text_content.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
                                else:
                                    text_content.append(f"--- Page {page_num + 1} ---\n[No text detected]")
                                    
                            except Exception as ocr_error:
                                text_content.append(f"--- Page {page_num + 1} ---\n[OCR failed: {str(ocr_error)}]")
                        else:
                            text_content.append(f"--- Page {page_num + 1} ---\n[No text content and OCR not available]")
            
            return self._clean_text("\n\n".join(text_content))
            
        except Exception as e:
            logging.error(f"PDF extraction failed for {pdf_path}: {e}")
            return f"[PDF EXTRACTION FAILED] Error extracting text from {pdf_path.name}: {str(e)}"
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = [line.strip() for line in text.split('\n')]
        cleaned_lines = [line for line in lines if line]
        
        # Join lines with proper spacing
        return '\n'.join(cleaned_lines)
    
    def extract_from_multiple_files(self, file_paths: List[Union[str, Path]]) -> List[dict]:
        """
        Extract text from multiple files
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            List of dictionaries with file info and extracted text
        """
        results = []
        
        for file_path in file_paths:
            file_path = Path(file_path)
            
            try:
                extracted_text = self.extract_text(file_path)
                results.append({
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "success": True,
                    "extracted_text": extracted_text,
                    "text_length": len(extracted_text)
                })
            except Exception as e:
                results.append({
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "success": False,
                    "error": str(e),
                    "extracted_text": "",
                    "text_length": 0
                })
        
        return results
    
    def is_supported_format(self, file_path: Union[str, Path]) -> bool:
        """
        Check if file format is supported
        
        Args:
            file_path: Path to file
            
        Returns:
            True if format is supported
        """
        file_extension = Path(file_path).suffix.lower()
        return file_extension in (
            self.supported_image_formats | 
            self.supported_pdf_formats | 
            self.supported_text_formats | 
            self.supported_doc_formats
        )

# Convenience function for quick text extraction
def extract_text_from_file(file_path: Union[str, Path]) -> str:
    """
    Quick text extraction from a single file
    
    Args:
        file_path: Path to file
        
    Returns:
        Extracted text
    """
    extractor = OCRExtractor()
    return extractor.extract_text(file_path)
